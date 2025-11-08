"""
FileInputAgent - Extracts text from PDF, DOCX, and TXT files.
"""
import os
import fitz  # PyMuPDF
from docx import Document
from typing import Union, BinaryIO
import logging

logger = logging.getLogger(__name__)


class FileInputAgent:
    """
    Agent responsible for extracting text from various file formats.
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt'}
    
    def extract_text(self, file_path: str = None, file_content: bytes = None, filename: str = None) -> str:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path: Path to the file
            file_content: File content as bytes (alternative to file_path)
            filename: Filename for determining file type
            
        Returns:
            Extracted text content
        """
        if file_path:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            filename = filename or os.path.basename(file_path)
            extension = os.path.splitext(filename)[1].lower()
            
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
        
        elif file_content and filename:
            extension = os.path.splitext(filename)[1].lower()
            
            if extension == '.pdf':
                return self._extract_from_pdf_bytes(file_content)
            elif extension in ['.docx', '.doc']:
                return self._extract_from_docx_bytes(file_content)
            elif extension == '.txt':
                return self._extract_from_txt_bytes(file_content)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
        
        else:
            raise ValueError("Either file_path or (file_content and filename) must be provided")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                text_parts.append(text)
            
            doc.close()
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise
    
    def _extract_from_pdf_bytes(self, file_content: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                text_parts.append(text)
            
            doc.close()
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF from bytes: {e}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise
    
    def _extract_from_docx_bytes(self, file_content: bytes) -> str:
        """Extract text from DOCX bytes."""
        try:
            from io import BytesIO
            doc = Document(BytesIO(file_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX from bytes: {e}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error extracting TXT: {e}")
            raise
    
    def _extract_from_txt_bytes(self, file_content: bytes) -> str:
        """Extract text from TXT bytes."""
        try:
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error extracting TXT from bytes: {e}")
            raise
    
    def is_supported(self, filename: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            filename: Name of the file
            
        Returns:
            True if supported, False otherwise
        """
        extension = os.path.splitext(filename)[1].lower()
        return extension in self.supported_formats

